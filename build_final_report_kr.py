from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from build_final_report import (
    BASE,
    RESULTS,
    set_cell_shading,
    set_cell_text,
    set_table_borders,
    add_page_number,
    add_heading,
    add_body,
    add_bullets,
    add_numbered,
    add_figure,
    load_csv,
    add_result_table,
    add_key_value_table,
)


OUT = BASE / "SPACE312_Final_Project_Report_KR_balanced.docx"


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    add_page_number(section)

    styles = doc.styles
    for name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Arial"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.size = Pt(24)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(13)


def add_cover(doc):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SPACE312 Final Project")
    r.bold = True
    r.font.size = Pt(25)
    r.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GTO-to-GEO Transfer Trajectory Optimization")
    r.bold = True
    r.font.size = Pt(17)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Mission Target: GEO-KOMPSAT-2A")
    r.font.size = Pt(13)

    doc.add_paragraph()
    add_key_value_table(
        doc,
        [
            ("교과목", "SPACE312 우주비행역학"),
            ("과제명", "Final Project"),
            ("학생", "2024105257 서혜인"),
            ("초기 시각", "2026-06-01 00:00:00 UTC"),
            ("동역학 모델", "지구 중심 2체 문제"),
            ("최적화 목표", "총 impulsive dV와 전이 시간의 Pareto 최적화"),
        ],
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run(
        "본 보고서는 최종 궤적 결과뿐 아니라, 왜 Lambert 전이만으로 끝내지 않고 "
        "GEO 반경 phasing 해를 추가로 탐색했는지를 논리적으로 정리한다."
    )
    r.italic = True
    r.font.size = Pt(10.5)
    doc.add_page_break()


def add_toc(doc):
    add_heading(doc, "목차", 1)
    items = [
        "1. 요약",
        "2. 문제 해석 및 기본 가정",
        "3. 목표 상태 검산",
        "4. 최적화 문제 정식화",
        "5. 풀이 전략: Lambert에서 phasing으로",
        "6. Baseline Lambert 전이 결과",
        "7. GEO 반경 phasing family",
        "8. Pareto 해석 및 대표 해",
        "9. 최종 선택 궤적과 가시성 분석",
        "10. 결론",
        "부록 A. MATLAB 구현 요약",
    ]
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)
    doc.add_page_break()


def build():
    rows = load_csv("FinalProject_ReportSolutions.csv")
    balanced = next((r for r in rows if r["Method"] == "Phasing"), rows[len(rows)//2])
    min_dv = rows[-1]
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_toc(doc)

    add_heading(doc, "1. 요약", 1)
    add_body(
        doc,
        "본 과제는 GTO에서 GEO-KOMPSAT-2A의 GEO 위치로 전이하는 궤적을 설계하는 문제이다. "
        "평가 기준은 하나의 최소값이 아니라 총 dV와 전이 시간 Delta t의 Pareto optimal set이므로, "
        "단순히 가장 빠른 궤적 또는 가장 dV가 작은 궤적 하나만 제시하는 것은 충분하지 않다고 판단하였다."
    )
    add_body(
        doc,
        "처음에는 강의에서 다룬 Lambert rendezvous 문제와 pork-chop 탐색 방법을 그대로 적용하였다. "
        "즉, 각 전이 시간에 대해 GEO target을 미래 시각까지 전파하고, 초기 GTO 위치에서 해당 target 위치까지 "
        "Lambert 문제를 풀어 dV를 계산하였다. 이 방법은 빠른 전이 영역의 Pareto 해를 잘 제공한다."
    )
    add_body(
        doc,
        "하지만 결과를 검토하는 과정에서 중요한 기하학적 특징을 발견하였다. 우주선은 GTO의 apogee에서 출발하며, "
        "이 지점의 반경은 이미 GEO 반경과 같다. 따라서 반경을 크게 바꾸는 전이가 아니라, 같은 GEO 반경 근처에서 "
        "위상을 맞추는 phasing orbit을 사용하면 더 작은 dV가 가능하다고 판단하였다. 이 판단에 따라 Lambert 해와 "
        "GEO 반경 phasing 해를 함께 탐색하여 최종 Pareto front를 구성하였다."
    )
    add_body(
        doc,
        f"최종적으로 보고서의 주 설계점은 균형 해로 선정하였다. 이 해는 Delta t = "
        f"{float(balanced['TOF_days']):.4f} days, dV1 = {float(balanced['dV1_km_s']):.4f} km/s, "
        f"dV2 = {float(balanced['dV2_km_s']):.4f} km/s, 총 dV = "
        f"{float(balanced['dVtotal_km_s']):.4f} km/s이다. 최저 dV 한계점은 "
        f"{float(min_dv['TOF_days']):.4f} days, {float(min_dv['dVtotal_km_s']):.4f} km/s로 별도 제시하였다."
    )

    add_heading(doc, "2. 문제 해석 및 기본 가정", 1)
    add_body(
        doc,
        "모든 위치와 속도는 ECI 좌표계에서 표현하였다. 전파 모델은 과제 조건에 맞추어 지구 중심 2체 문제만 사용하였다. "
        "따라서 J2, 달/태양 섭동, 대기저항 등은 고려하지 않았다. 초기 상태는 주어진 GTO apogee 상태이고, 목표 상태는 "
        "GEO-KOMPSAT-2A의 단순화된 동경 128.2 deg East에 해당하는 원형 GEO 상태이다."
    )
    add_key_value_table(
        doc,
        [
            ("mu_E", "398600.4418 km^3/s^2"),
            ("R_E", "6378.137 km"),
            ("omega_E", "7.2921150e-5 rad/s"),
            ("r_GEO", "42164.173 km"),
            ("초기 궤도", "GTO, inclination = 24.5 deg, apogee at t0"),
            ("목표 궤도", "Circular GEO, inclination = 0 deg"),
            ("전이 시간 제한", "1 day <= Delta t <= 30 days"),
        ],
    )
    add_body(
        doc,
        "중요한 점은 GEO target을 ECI에서 고정점으로 취급하면 안 된다는 것이다. GEO 위성은 지구 자전과 같은 각속도로 "
        "움직이므로, arrival time이 달라질 때마다 target의 ECI 위치도 달라진다. 따라서 각 후보 전이 시간마다 "
        "target state를 새로 계산하였다."
    )

    add_heading(doc, "3. 목표 상태 검산", 1)
    add_body(
        doc,
        "PDF의 표를 그대로 읽는 것보다, 주어진 coordinate assumption을 먼저 검산하는 것이 필요하다고 보았다. "
        "목표 GEO의 ECI 각도는 theta_ERA,0 + lambda_GEO로 계산된다. 이를 사용하면 t0에서의 목표 위치와 속도는 "
        "다음과 같이 얻어진다."
    )
    add_bullets(
        doc,
        [
            "r_GEO(t0) = [40244.6079, 12577.3213, 0] km",
            "v_GEO(t0) = [-0.917153, 2.934683, 0] km/s",
            "||r_GEO|| = 42164.173 km, ||v_GEO|| = 3.074660 km/s",
        ],
    )
    add_body(
        doc,
        "텍스트 추출 과정에서는 x 성분이 41244.6079처럼 보일 수 있지만, 이 값을 사용하면 반경이 GEO 반경이 아니고 "
        "주어진 속도와도 수직 조건이 맞지 않는다. 따라서 본 해석에서는 표의 숫자를 무조건 입력하기보다, PDF에 명시된 "
        "ERA와 GEO longitude 정의에서 목표 상태를 계산하도록 하였다. 이 선택이 물리적으로 더 일관적이다."
    )

    add_heading(doc, "4. 최적화 문제 정식화", 1)
    add_body(
        doc,
        "목적함수는 F(x) = [sum dV, Delta t]로 두었다. 어떤 해가 Pareto optimal이라는 것은, 다른 유효 해가 "
        "동시에 더 작은 dV와 더 짧은 전이 시간을 갖지 못한다는 뜻이다. 따라서 빠르지만 dV가 큰 해와 느리지만 "
        "dV가 작은 해가 동시에 의미 있는 해가 될 수 있다."
    )
    add_body(doc, "두 impulsive maneuver에 대한 dV 계산은 다음과 같이 하였다.")
    add_bullets(
        doc,
        [
            "dV1 = ||v1_plus - v_GTO(t0)||",
            "dV2 = ||v_GEO(tf) - v2_minus||",
            "dVtotal = dV1 + dV2",
        ],
    )

    add_heading(doc, "5. 풀이 전략: Lambert에서 phasing으로", 1)
    add_body(
        doc,
        "처음부터 phasing 해를 가정한 것은 아니다. 강의 코드와 Homework 3의 흐름을 보면, 교수님이 기대하는 기본 접근은 "
        "Lambert solver를 이용한 boundary-value transfer와 pork-chop 형태의 탐색이라고 판단하였다. 그래서 먼저 "
        "이 baseline을 구현하였다."
    )
    add_numbered(
        doc,
        [
            "1~30일 범위에서 전이 시간을 조밀하게 sweep한다.",
            "각 Delta t마다 GEO-KOMPSAT-2A의 미래 ECI 위치와 속도를 계산한다.",
            "초기 GTO 위치에서 미래 target 위치까지 Lambert 문제를 푼다.",
            "초기 burn과 최종 insertion burn을 계산한다.",
            "모든 후보를 Pareto dominance 기준으로 필터링한다.",
        ],
    )
    add_body(
        doc,
        "그런데 Lambert 결과만 보면 짧은 시간에서는 dV가 빠르게 줄어들지만, 장시간 영역에서는 더 좋은 구조적 해를 "
        "놓칠 수 있다고 보았다. 특히 초기 위치가 GEO 반경이라는 사실은 단순한 우연이 아니라 궤적 설계에서 활용할 수 "
        "있는 조건이다. 이 때문에 두 번째 탐색 family로 same-point phasing orbit을 추가하였다."
    )

    add_heading(doc, "6. Baseline Lambert 전이 결과", 1)
    add_body(
        doc,
        "Lambert 해는 빠른 전이 영역을 담당한다. 예를 들어 1일 전이 해는 요구되는 시간이 가장 짧지만 총 dV가 크다. "
        "전이 시간을 조금 늘리면 target의 상대 위치가 달라져 dV가 감소하지만, Lambert arc만으로는 장시간 저에너지 "
        "phasing 해만큼 dV를 낮추기 어렵다."
    )
    add_figure(doc, "Pareto_dV_vs_TOF.png", "그림 1. Lambert 및 phasing 후보를 포함한 Pareto front.", 6.2)

    add_heading(doc, "7. GEO 반경 phasing family", 1)
    add_body(
        doc,
        "phasing 해의 핵심 아이디어는 다음과 같다. 우주선은 이미 GEO 반경의 GTO apogee에 있으므로, 출발점 자체가 "
        "GEO 궤도 위의 한 점이다. GEO-KOMPSAT-2A가 시간이 지나 이 같은 ECI 방향으로 돌아오는 시각을 arrival time으로 "
        "선택하면, 우주선은 그동안 정수 회전 phasing orbit을 돌고 같은 점에 돌아와 rendezvous할 수 있다."
    )
    add_body(
        doc,
        "이 해는 과제 조건을 벗어난 것이 아니다. 여전히 지구 중심 2체 전파이고, impulsive maneuver만 사용하며, "
        "최종 상태는 움직이는 GEO target에 맞춘다. 다만 Lambert의 일반 boundary-value 해를 무작정 찾는 대신, "
        "문제의 특수한 초기 반경 조건을 이용하여 더 좋은 후보군을 구성한 것이다."
    )
    add_body(
        doc,
        "전이 시간이 길어질수록 phasing orbit의 주기는 GEO 주기에 가까워지고, 따라서 최종 GEO insertion dV는 매우 작아진다. "
        "반면 초기 dV는 GTO의 경사 성분을 제거하고 equatorial phasing orbit으로 들어가기 위해 여전히 필요하다. "
        "이 때문에 phasing family의 총 dV는 약 1.75 km/s 근처에서 천천히 감소한다."
    )

    add_heading(doc, "8. Pareto 해석 및 대표 해", 1)
    add_result_table(doc, rows)
    add_body(
        doc,
        "표에서 볼 수 있듯이, 1일 Lambert 해는 빠르지만 dV가 크고, 29.37일 phasing 해는 가장 작은 dV를 갖지만 시간이 길다. "
        "따라서 어느 하나가 다른 하나를 완전히 지배하지 않는다. 이것이 Pareto front의 의미이다. 실제 임무에서는 연료를 "
        "우선할지, 임무 투입 시간을 우선할지에 따라 다른 design point를 선택해야 한다."
    )
    add_body(
        doc,
        "특히 3.44일 phasing 해는 총 dV가 1.7611 km/s로 크게 낮으면서도 30일 전체를 사용하지 않는다. "
        "29일 해와 비교하면 dV 증가는 약 0.0120 km/s에 불과하지만 전이 시간은 약 25.93일 줄어든다. "
        "따라서 본 보고서의 주 설계점은 3.44일 phasing 해로 두고, 1일 해와 29일 해는 각각 빠른 전이 한계와 "
        "최저 dV 한계로 구분하여 제시하였다."
    )

    add_heading(doc, "9. 최종 선택 궤적과 가시성 분석", 1)
    add_figure(doc, "Trajectory3D.png", "그림 2. 선택한 전이 궤적, 초기 GTO, 이동하는 GEO target.", 6.0)
    add_figure(doc, "Radius_vs_Time.png", "그림 3. 선택 궤적의 radius history.", 5.8)
    add_figure(doc, "DeltaV_Maneuvers.png", "그림 4. 초기 burn과 최종 GEO insertion burn.", 5.2)
    add_figure(doc, "Position_Error_Final_Window.png", "그림 5. 도착 직전 moving target에 대한 position error.", 5.8)
    add_figure(doc, "KHU_Visibility.png", "그림 6. KHU 기준 elevation 및 visibility interval.", 5.8)
    add_body(
        doc,
        "최종 그림들은 균형 설계점인 3.44일 phasing 해를 기준으로 작성하였다. 선택한 궤적이 고정된 GEO 점이 아니라 시간에 따라 움직이는 target을 향해 수렴함을 확인하기 위해 작성하였다. "
        "가시성은 KHU 위치에서 ECI 상태를 ECEF로 변환한 뒤 SEZ 좌표계의 elevation이 0도보다 큰 구간으로 정의하였다."
    )

    add_heading(doc, "10. 결론", 1)
    add_body(
        doc,
        "본 프로젝트에서 가장 중요한 판단은 Lambert 해를 baseline으로 사용하되, 그 결과에 만족하지 않고 초기 조건의 기하학적 "
        "의미를 다시 본 것이다. 초기 GTO apogee가 GEO 반경에 있다는 점은 저에너지 phasing transfer를 가능하게 한다. "
        "따라서 최종 Pareto set은 짧은 시간의 Lambert 해와 낮은 dV의 phasing 해를 함께 포함해야 한다."
    )
    add_body(
        doc,
        f"주 설계점으로 추천하는 균형 해는 Delta t = {float(balanced['TOF_days']):.4f} days, "
        f"dV1 = {float(balanced['dV1_km_s']):.4f} km/s, dV2 = {float(balanced['dV2_km_s']):.4f} km/s, "
        f"dVtotal = {float(balanced['dVtotal_km_s']):.4f} km/s이다. 최저 dV 후보는 "
        f"Delta t = {float(min_dv['TOF_days']):.4f} days, dVtotal = {float(min_dv['dVtotal_km_s']):.4f} km/s이지만, "
        "전이 시간이 지나치게 길기 때문에 mission design point라기보다는 low-energy limit으로 해석하였다. "
        "반대로 빠른 투입이 중요하다면 1~1.3일 Lambert 해를 선택할 수 있다. 이와 같이 본 결과는 과제에서 요구한 "
        "dV와 Delta t의 trade space를 직접 보여준다."
    )

    add_heading(doc, "부록 A. MATLAB 구현 요약", 1)
    add_bullets(
        doc,
        [
            "SPACE312_Final_Project_Main.m은 외부 파일 없이 실행되도록 local function을 포함한다.",
            "LambertTd 함수는 강의 코드의 universal variable Lambert 구조를 바탕으로 하였다.",
            "target GEO 상태는 ERA와 longitude로 매 시각 계산한다.",
            "Lambert 후보와 phasing 후보를 모두 생성한 뒤 Pareto dominance로 필터링한다.",
            "결과 CSV, 3D trajectory, radius history, dV plot, final error plot, KHU visibility plot, GIF animation을 results 폴더에 저장한다.",
        ],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
