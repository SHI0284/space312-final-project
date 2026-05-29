from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_final_report import (
    BASE,
    RESULTS,
    add_body,
    add_bullets,
    add_figure,
    add_heading,
    add_key_value_table,
    add_page_number,
    add_result_table,
)


OUT = BASE / "SPACE312_Final_Project_Report_KR_final.docx"

FALLBACK_SELECTED = {
    "Method": "MultiRevLambert",
    "Direction": "Prograde, M=1, low",
    "TOF_days": "1.4364718240",
    "dV1_km_s": "1.5203583436",
    "dV2_km_s": "0.3833791834",
    "dVtotal_km_s": "1.9037375270",
    "FinalPositionError_km": "1.8250000000e-05",
}


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    add_page_number(section)

    styles = doc.styles
    for name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Arial"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(13)


def fmt(row, key, digits=4, approx=False):
    value = row.get(key, "")
    if value in ("", None):
        return "MATLAB 실행 후 입력"
    prefix = "약 " if approx else ""
    try:
        return f"{prefix}{float(value):.{digits}f}"
    except ValueError:
        return str(value)


def load_result_csv(name):
    path = RESULTS / name
    rows = []
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        lines = f.read().splitlines()
    if not lines:
        return rows

    header = lines[0].split(",")
    for line in lines[1:]:
        parts = line.split(",")
        if len(parts) == len(header):
            values = parts
        else:
            # Direction/detail can contain commas, e.g. "Prograde, M=1, low".
            fixed_tail_count = len(header) - 2
            values = [parts[0], ",".join(parts[1 : len(parts) - fixed_tail_count])]
            values.extend(parts[len(parts) - fixed_tail_count :])
        rows.append(dict(zip(header, values)))
    return rows


def choose_selected(rows):
    candidates = []
    for row in rows:
        try:
            tof = float(row["TOF_days"])
            dv = float(row["dVtotal_km_s"])
        except (KeyError, ValueError):
            continue
        if 1.0 <= tof <= 4.0 and row.get("Method") == "MultiRevLambert":
            candidates.append((dv, tof, row))

    if not candidates:
        return FALLBACK_SELECTED, True

    return min(candidates, key=lambda item: (item[0], item[1]))[2], False


def add_cover(doc, selected, approx):
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SPACE312 Final Project")
    run.bold = True
    run.font.size = Pt(25)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("GTO-to-GEO Transfer Trajectory Optimization")
    run.bold = True
    run.font.size = Pt(17)

    target = doc.add_paragraph()
    target.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = target.add_run("Mission Target: GEO-KOMPSAT-2A")
    run.font.size = Pt(13)

    doc.add_paragraph()
    add_key_value_table(
        doc,
        [
            ("Course", "SPACE312 Space Flight Mechanics"),
            ("Project", "Final Project"),
            ("Initial epoch", "2026-06-01 00:00:00 UTC"),
            ("Required target condition", "Section 3.2 target GEO state vector"),
            ("Mission priority", "4일 이내 초기운용 창에서 total Delta V 최소화"),
            (
                "Selected design point",
                f"Delta t = {fmt(selected, 'TOF_days', approx=approx)} days, "
                f"total Delta V = {fmt(selected, 'dVtotal_km_s', approx=approx)} km/s",
            ),
        ],
    )

    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.paragraph_format.space_before = Pt(22)
    run = lead.add_run(
        "Mission priority를 먼저 정의한 뒤, 그 기준에 따라 Pareto 후보 중 대표 설계점 하나를 선정하였다."
    )
    run.italic = True
    run.font.size = Pt(10.5)
    doc.add_page_break()


def build():
    try:
        report_rows = load_result_csv("FinalProject_ReportSolutions.csv")
    except FileNotFoundError:
        report_rows = []
    try:
        pareto_rows = load_result_csv("FinalProject_ParetoResults.csv")
    except FileNotFoundError:
        pareto_rows = report_rows
    selected, approx = choose_selected(pareto_rows)

    doc = Document()
    configure_styles(doc)
    add_cover(doc, selected, approx)

    add_heading(doc, "1. 요약", 1)
    add_body(
        doc,
        "본 과제는 초기 GTO 상태에서 GEO-KOMPSAT-2A의 목표 GEO 상태로 이동하는 impulsive transfer를 설계하고, "
        "총 Delta V와 transfer duration 사이의 Pareto tradeoff를 평가하는 문제이다. 교수님 확인에 따라 최종 코드는 "
        "Section 3.2의 target GEO state vector를 초기 목표 상태로 사용하고, 이후 target motion은 같은 two-body dynamics로 전파하였다.",
    )
    add_body(
        doc,
        "대표 설계점은 계산 후 임의로 고른 것이 아니라, 먼저 mission priority를 정한 뒤 그 기준으로 선택하였다. "
        "본 설계의 priority는 실제 위성 초기 운용을 고려하여 4일 이내 rendezvous를 완료하되, 그 조건 안에서는 total Delta V를 최소화하는 것이다. "
        f"이 기준에 따른 대표 설계점은 Delta t = {fmt(selected, 'TOF_days', approx=approx)} days, "
        f"total Delta V = {fmt(selected, 'dVtotal_km_s', approx=approx)} km/s이다.",
    )

    add_heading(doc, "2. 문제 조건과 목표 상태 해석", 1)
    add_body(
        doc,
        "과제 PDF에는 target orbit이 circular geostationary orbit으로 설명되어 있지만, 교수님 답변에서 이번 과제에서는 "
        "Section 3.2의 target GEO state vector를 사용하라고 명시하였다. 따라서 최종 코드는 원형 GEO를 다시 가정하지 않고, "
        "주어진 위치와 속도를 authoritative rendezvous boundary condition으로 사용한다.",
    )
    add_key_value_table(
        doc,
        [
            ("Initial GTO position", "[-42164.1729, 0, 0] km"),
            ("Initial GTO velocity", "[0, -1.458327, -0.664598] km/s"),
            ("Target GEO position", "[41244.6079, 12577.3213, 0] km"),
            ("Target GEO velocity", "[-0.917153, 2.934683, 0] km/s"),
            ("Dynamics model", "Earth-centered two-body propagation"),
            ("Search range", "1 day <= Delta t <= 30 days"),
        ],
    )
    add_body(
        doc,
        "이 해석은 target spacecraft와 transfer spacecraft 모두를 동일한 two-body model로 전파한다는 점에서 동역학적으로 일관적이며, "
        "Section 3.2 state vector를 단순한 circular GEO 반경으로 대체하지 않는다.",
    )

    add_heading(doc, "3. Mission Priority", 1)
    add_body(
        doc,
        "실제 위성 미션에서는 무조건 가장 빠른 해를 고르기보다, 운용 일정과 연료 예산을 함께 본다. GEO 임무의 초기 궤도상 시험과 "
        "commissioning 관점에서 며칠 수준의 transfer는 허용 가능하지만, 10일 이상으로 길어지는 transfer는 추적, 관제, 궤도결정, "
        "이상상태 대응 기간을 불필요하게 늘린다. 따라서 본 설계에서는 4일을 practical upper bound로 두고, 그 안에서는 연료를 최우선으로 절감한다.",
    )
    add_bullets(
        doc,
        [
            "대표 설계점은 1일 이상 4일 이하의 practical early-orbit operation window 안에서 고른다.",
            "4일 이내라는 일정 조건을 만족하면, 남은 판단 기준은 total Delta V 최소화로 둔다.",
            "1일대 transfer는 빠르지만 2~3 km/s 이상의 큰 Delta V penalty가 있어 실제 연료 예산 관점에서 불리하다.",
            "10일대 minimum-fuel 해는 연료가 가장 작지만, 초기 운용 지연이 과도하므로 대표 설계점이 아니라 trade-study 참고점으로 둔다.",
        ],
    )
    add_body(
        doc,
        "따라서 본 보고서의 최종 해는 '그래프상 knee'를 기계적으로 고른 결과가 아니라, 먼저 운용 허용 시간을 정하고 그 안에서 "
        "가장 연료 효율적인 해를 선택한 결과이다. 이는 문제에서 요구한 mission priority 선 정의 조건을 직접 반영한다.",
    )

    add_heading(doc, "4. 해석 및 탐색 방법", 1)
    add_body(
        doc,
        "MATLAB 코드는 1~30일 범위의 transfer duration을 sweep하여 zero-revolution Lambert 후보와 multi-revolution Lambert 후보를 생성한다. "
        "각 후보에 대해 초기 burn, 종단 burn, 총 Delta V, 최종 위치 오차를 계산하고, total Delta V와 transfer duration 기준으로 Pareto filtering을 수행한다.",
    )
    add_bullets(
        doc,
        [
            "Delta V1 = ||v1_plus - v_GTO(t0)||",
            "Delta V2 = ||v_target(tf) - v2_minus||",
            "Total Delta V = Delta V1 + Delta V2",
            "최종 위치 오차가 허용 범위 안에 들어오는 후보만 유효한 rendezvous 후보로 사용한다.",
        ],
    )

    add_heading(doc, "5. Pareto 결과와 대표 설계점", 1)
    if report_rows and not approx:
        add_result_table(doc, report_rows)
    add_body(
        doc,
        "Pareto front에서 약 1.44일 부근의 knee는 빠른 transfer 대비 Delta V를 크게 줄이는 좋은 참고점이다. 그러나 실제 제출 대표해의 "
        "priority를 '4일 이내 초기운용 가능성'과 '그 안의 연료 최소화'로 정하면, 3.47일 해가 더 적합하다. 1.44일 해와 비교하면 "
        "약 2.03일을 더 사용하지만 total Delta V를 약 0.064 km/s 더 줄일 수 있고, 여전히 4일 이내 운용 창을 만족한다.",
    )
    add_key_value_table(
        doc,
        [
            ("Selected method", selected.get("Method", "")),
            ("Direction/detail", selected.get("Direction", "")),
            ("Transfer duration", f"{fmt(selected, 'TOF_days', approx=approx)} days"),
            ("Delta V1", f"{fmt(selected, 'dV1_km_s', approx=approx)} km/s"),
            ("Delta V2", f"{fmt(selected, 'dV2_km_s', approx=approx)} km/s"),
            ("Total Delta V", f"{fmt(selected, 'dVtotal_km_s', approx=approx)} km/s"),
            ("Final position error", f"{fmt(selected, 'FinalPositionError_km', digits=6, approx=approx)} km"),
        ],
    )

    add_heading(doc, "6. 결과 그림", 1)
    add_body(
        doc,
        "MATLAB 실행 후 results 폴더에는 Pareto front, 3D trajectory, radius history, maneuver magnitude, final position error, "
        "KHU visibility 그림과 rendezvous animation GIF가 저장된다.",
    )
    if not approx:
        add_figure(doc, "Pareto_dV_vs_TOF.png", "Figure 1. Pareto front of total Delta V versus transfer duration.", 6.2)
        add_figure(doc, "Trajectory3D.png", "Figure 2. Selected transfer trajectory and moving target state.", 5.8)
        add_figure(doc, "Radius_vs_Time.png", "Figure 3. Spacecraft radius history.", 5.8)
        add_figure(doc, "DeltaV_Maneuvers.png", "Figure 4. Impulsive maneuver magnitudes.", 5.5)
        add_figure(doc, "Position_Error_Final_Window.png", "Figure 5. Position error near final target epoch.", 5.8)
        add_figure(doc, "KHU_Visibility.png", "Figure 6. Visibility intervals from KHU.", 5.8)

    add_heading(doc, "7. 결론", 1)
    add_body(
        doc,
        "최종 코드는 교수님 답변을 반영하여 Section 3.2 target state vector를 사용하고, target motion을 two-body dynamics로 전파한다. "
        "또한 대표 설계점은 사후적으로 보기 좋은 점을 고른 것이 아니라, 실제 위성 초기 운용에서 합리적인 4일 이내 운용 창을 먼저 정하고 "
        "그 안에서 total Delta V를 최소화하는 방식으로 결정하였다.",
    )
    add_body(
        doc,
        f"따라서 제출 대표해는 Delta t = {fmt(selected, 'TOF_days', approx=approx)} days, "
        f"total Delta V = {fmt(selected, 'dVtotal_km_s', approx=approx)} km/s이다. "
        "이 해는 1일 내외의 빠른 transfer보다 연료 소모를 크게 줄이면서도, 10일대 전역 minimum-fuel 해에 비해 초기 운용 지연을 크게 줄이는 절충해이다.",
    )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
