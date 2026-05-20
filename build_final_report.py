from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import csv


BASE = Path(__file__).resolve().parent
RESULTS = BASE / "results"
OUT = BASE / "SPACE312_Final_Project_Report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="B7C4D6", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.runs[0].font.color.rgb = RGBColor(31, 78, 121)
    elif level == 2:
        p.runs[0].font.color.rgb = RGBColor(46, 97, 142)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.12
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(91, 103, 112)


def add_figure(doc, filename, caption, width=6.0):
    path = RESULTS / filename
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        add_caption(doc, caption)


def load_csv(name):
    path = RESULTS / name
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        raw = list(csv.reader(f))
    header = raw[0]
    rows = []
    for values in raw[1:]:
        if len(values) == len(header) + 1 and values[0] == "Phasing":
            values = [values[0], values[1] + ", " + values[2]] + values[3:]
        rows.append(dict(zip(header, values)))
    return rows


def add_result_table(doc, rows):
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    headers = ["Case", "Method", "TOF [days]", "dV1 [km/s]", "dV2 [km/s]", "Total dV [km/s]"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=(255, 255, 255))
        set_cell_shading(table.rows[0].cells[i], "1F4E79")

    for idx, row in enumerate(rows, 1):
        cells = table.add_row().cells
        values = [
            idx,
            row["Method"],
            f'{float(row["TOF_days"]):.4f}',
            f'{float(row["dV1_km_s"]):.4f}',
            f'{float(row["dV2_km_s"]):.4f}',
            f'{float(row["dVtotal_km_s"]):.4f}',
        ]
        for j, value in enumerate(values):
            set_cell_text(cells[j], value)
            if idx % 2 == 0:
                set_cell_shading(cells[j], "EEF3F8")
    doc.add_paragraph()


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, bold=True)
        set_cell_text(cells[1], value)
        set_cell_shading(cells[0], "D9EAF7")
    doc.add_paragraph()


def build():
    report_rows = load_csv("FinalProject_ReportSolutions.csv")
    best = report_rows[-1]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    add_page_number(section)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Arial"
    styles["Title"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    styles["Title"].font.size = Pt(24)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    styles["Heading 2"].font.size = Pt(13)

    # Cover
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("SPACE312 Final Project")
    r.bold = True
    r.font.size = Pt(25)
    r.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("GTO-to-GEO Transfer Trajectory Optimization")
    r.bold = True
    r.font.size = Pt(17)

    target = doc.add_paragraph()
    target.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = target.add_run("Mission Target: GEO-KOMPSAT-2A")
    r.font.size = Pt(13)

    doc.add_paragraph()
    add_key_value_table(
        doc,
        [
            ("Course", "SPACE312 Space Flight Mechanics"),
            ("Project", "Final Project"),
            ("Student", "2024105257 서혜인"),
            ("Initial epoch", "2026-06-01 00:00:00 UTC"),
            ("Dynamics model", "Two-body Earth-centered inertial propagation"),
            ("Submission focus", "Pareto-optimal total dV versus transfer duration"),
        ],
    )

    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.paragraph_format.space_before = Pt(24)
    rr = lead.add_run(
        "This report explains not only the final trajectory results, but also the reasoning path that led "
        "from a basic Lambert transfer search to a lower-energy phasing-orbit family."
    )
    rr.italic = True
    rr.font.size = Pt(10.5)

    doc.add_page_break()

    # Static table of contents
    add_heading(doc, "Table of Contents", 1)
    toc_items = [
        "1. Executive Summary",
        "2. Mission Interpretation and Key Assumptions",
        "3. Optimization Formulation",
        "4. Search Strategy and Reasoning Process",
        "5. Baseline Lambert Transfer Results",
        "6. GEO-Radius Phasing Family",
        "7. Pareto Discussion",
        "8. Selected Trajectory and Visibility Results",
        "9. Conclusion",
        "Appendix A. Implementation Notes",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)

    doc.add_page_break()

    add_heading(doc, "1. Executive Summary", 1)
    add_body(
        doc,
        "The final project was interpreted as a two-objective trajectory optimization problem: minimize the "
        "total impulsive velocity increment while also minimizing the transfer duration. A single 'best' "
        "trajectory is therefore not sufficient. Instead, the solution should present a Pareto set that shows "
        "the tradeoff between fast transfers and low-energy transfers."
    )
    add_body(
        doc,
        "My first design choice was to reproduce the method emphasized in the lecture examples: use Lambert's "
        "problem for rendezvous targeting, propagate the moving GEO target to the selected arrival time, and "
        "sweep the time of flight in a pork-chop style search. This produced a valid baseline Pareto front. "
        "However, after inspecting the mission geometry, I noticed that the spacecraft begins at GTO apogee, "
        "whose radius is already the GEO radius. This suggested an additional family of solutions: leave the "
        "initial point onto a nearly GEO-radius equatorial phasing orbit, wait until GEO-KOMPSAT-2A returns to "
        "the same inertial direction, and then perform a small GEO insertion burn."
    )
    add_body(
        doc,
        f"The resulting Pareto set contains both fast Lambert transfers and low-dV phasing transfers. The "
        f"lowest-dV representative solution found in the 1 to 30 day window has transfer duration "
        f"{float(best['TOF_days']):.4f} days and total dV {float(best['dVtotal_km_s']):.4f} km/s."
    )

    add_heading(doc, "2. Mission Interpretation and Key Assumptions", 1)
    add_body(
        doc,
        "All propagation was performed in the Earth-centered inertial frame under the two-body model. The "
        "initial spacecraft state is the given GTO apogee state at the project epoch, and the target state is "
        "the circular GEO location corresponding to GEO-KOMPSAT-2A's simplified longitude. The target is not "
        "held fixed in ECI; it moves with the GEO angular rate during the transfer."
    )
    add_key_value_table(
        doc,
        [
            ("Earth gravitational parameter", "mu_E = 398600.4418 km^3/s^2"),
            ("Earth equatorial radius", "R_E = 6378.137 km"),
            ("Earth rotation rate", "omega_E = 7.2921150e-5 rad/s"),
            ("GEO radius", "r_GEO = 42164.173 km"),
            ("GEO longitude", "lambda_GEO = 128.2 deg East"),
            ("Transfer duration bound", "1 day <= Delta t <= 30 days"),
        ],
    )
    add_body(
        doc,
        "One important consistency check was applied to the target state. The coordinate assumptions define the "
        "target ECI angle as theta_ERA,0 + lambda_GEO. Using this relation gives r_GEO(t0) = "
        "[40244.6079, 12577.3213, 0] km and v_GEO(t0) = [-0.917153, 2.934683, 0] km/s. This is consistent "
        "with the GEO radius and circular speed. The value 41244.6079 sometimes appears from PDF text "
        "extraction, but that value would not have the GEO radius and would not be perpendicular to the listed "
        "velocity. Therefore, the code computes the target state from the stated coordinate assumptions rather "
        "than relying on a possibly misread table entry."
    )

    add_heading(doc, "3. Optimization Formulation", 1)
    add_body(
        doc,
        "The objective vector is F(x) = [sum dV, Delta t]. A candidate solution is Pareto-optimal if no other "
        "valid candidate has both smaller or equal total dV and smaller or equal transfer duration, with at "
        "least one of the two objectives strictly better."
    )
    add_body(
        doc,
        "For a two-impulse Lambert candidate, the departure and arrival impulses were computed as"
    )
    add_bullets(
        doc,
        [
            "dV1 = ||v1_plus - v_GTO(t0)||",
            "dV2 = ||v_GEO(tf) - v2_minus||",
            "dVtotal = dV1 + dV2",
        ],
    )
    add_body(
        doc,
        "For a phasing candidate, the same impulse accounting was used, but the transfer orbit is constrained "
        "to begin and end at the same GEO-radius inertial point after an integer number of revolutions. This "
        "keeps the comparison fair: both families are judged by the same objective vector."
    )

    add_heading(doc, "4. Search Strategy and Reasoning Process", 1)
    add_body(
        doc,
        "The search process was deliberately built in two layers. The first layer follows the lecture method "
        "directly: solve Lambert's problem for many transfer durations. This is the most natural starting point "
        "because the project is a boundary-value rendezvous problem and the course examples used Lambert "
        "targeting and pork-chop searches for similar decisions."
    )
    add_numbered(
        doc,
        [
            "Propagate the moving GEO target to each candidate arrival time.",
            "Solve Lambert's problem from the initial GTO apogee position to that future target position.",
            "Compute departure and insertion dV from the initial GTO and final GEO reference velocities.",
            "Filter all candidates using Pareto dominance.",
            "Inspect the geometry of the best and worst regions of the Pareto set.",
        ],
    )
    add_body(
        doc,
        "The fifth step is where the main improvement came from. The initial point is not an arbitrary GTO "
        "position. It is the apogee, and its radius is essentially the GEO radius. That means a low-energy "
        "solution does not necessarily need to travel from one radius to another. Instead, it can use time as "
        "a resource: adjust the orbital period so the spacecraft revisits the same inertial point when the "
        "GEO target has also rotated to that point."
    )

    add_heading(doc, "5. Baseline Lambert Transfer Results", 1)
    add_body(
        doc,
        "The Lambert sweep provides the fast-transfer side of the Pareto front. These solutions are useful "
        "because they satisfy the moving-target boundary condition with short flight times. Their drawback is "
        "that a short time of flight forces a relatively energetic transfer and a larger final insertion burn."
    )
    add_figure(doc, "Pareto_dV_vs_TOF.png", "Figure 1. Pareto front combining Lambert and phasing candidates.", 6.2)
    add_body(
        doc,
        "The first few Pareto points are Lambert solutions. For example, the 1-day solution is very fast but "
        "requires more than 3 km/s total dV. Increasing the time of flight to about 1.3 days reduces the total "
        "dV substantially, but Lambert solutions alone still do not exploit the full low-energy opportunity "
        "created by the initial GEO-radius apogee."
    )

    add_heading(doc, "6. GEO-Radius Phasing Family", 1)
    add_body(
        doc,
        "The phasing family was added because it follows directly from the geometry. If GEO-KOMPSAT-2A reaches "
        "the initial apogee inertial direction at the final time, then the spacecraft can depart from that same "
        "point, complete an integer number of revolutions on a phasing orbit, and return to the target point. "
        "The final burn then only circularizes or adjusts the small remaining speed difference."
    )
    add_body(
        doc,
        "For each possible arrival time satisfying the target phase condition, I tested integer revolution "
        "counts and selected physically valid phasing orbits whose perigee stayed above the Earth. This is not "
        "a separate objective or an unfair shortcut; it is simply a structured subset of valid two-body, "
        "impulsive trajectories. It is included because the project asks for optimized trajectories, not only "
        "Lambert arcs."
    )
    add_body(
        doc,
        "The best phasing solutions have large transfer durations but low dV. The departure burn mainly removes "
        "the initial GTO inclination component and changes the speed into an equatorial phasing orbit. As the "
        "duration increases, the phasing orbit period becomes closer to one sidereal day, so the final insertion "
        "burn becomes very small."
    )

    add_heading(doc, "7. Pareto Discussion", 1)
    add_result_table(doc, report_rows)
    add_body(
        doc,
        "The table shows why the answer should not be reported as one trajectory only. The 1-day Lambert "
        "solution is the best if schedule is the dominant requirement. The 29.37-day phasing solution is the "
        "best if propellant is the dominant requirement. Intermediate points, especially the 3.44-day phasing "
        "solution, are useful compromises because most of the dV reduction is achieved without using the full "
        "30-day limit."
    )
    add_body(
        doc,
        "The final selected trajectory for detailed plotting is the minimum-dV Pareto point. This is not the "
        "only acceptable design point, but it is the strongest demonstration that the trade space was searched "
        "beyond the basic Lambert baseline."
    )

    add_heading(doc, "8. Selected Trajectory and Visibility Results", 1)
    add_figure(doc, "Trajectory3D.png", "Figure 2. 3D view of selected transfer, initial GTO, and moving GEO target.", 6.0)
    add_figure(doc, "Radius_vs_Time.png", "Figure 3. Spacecraft radius history for the selected trajectory.", 5.8)
    add_figure(doc, "DeltaV_Maneuvers.png", "Figure 4. Impulsive maneuver magnitudes for the selected design point.", 5.2)
    add_figure(doc, "Position_Error_Final_Window.png", "Figure 5. Position error to the moving GEO target near arrival.", 5.8)
    add_figure(doc, "KHU_Visibility.png", "Figure 6. Visibility intervals from KHU for the selected trajectory.", 5.8)
    add_body(
        doc,
        "The selected trajectory reaches the target position with numerical terminal error on the order of "
        "meters or less in the search result. The plotted final-window error confirms that the spacecraft "
        "approaches the moving GEO target rather than a fixed point in ECI. Visibility from KHU was computed "
        "by transforming the spacecraft ECI position to ECEF using the Earth rotation angle, then converting "
        "the relative vector to the local SEZ frame and checking elevation above the horizon."
    )

    add_heading(doc, "9. Conclusion", 1)
    add_body(
        doc,
        "The final Pareto set supports the expected mission tradeoff: shorter transfers require larger total "
        "dV, while longer transfers allow the spacecraft to use orbital phasing and reduce propellant cost. "
        "The most important design decision was to not stop at the first Lambert sweep. Lambert targeting was "
        "necessary to establish the baseline and to match the lecture methodology, but the mission geometry "
        "itself suggested a better low-energy family. Because the spacecraft begins at GEO radius, the optimized "
        "solution can trade time for propellant by waiting for GEO-KOMPSAT-2A to return to the same inertial "
        "direction."
    )
    add_body(
        doc,
        f"Within the project constraints, the minimum-dV candidate found has Delta t = "
        f"{float(best['TOF_days']):.4f} days, dV1 = {float(best['dV1_km_s']):.4f} km/s, "
        f"dV2 = {float(best['dV2_km_s']):.4f} km/s, and total dV = "
        f"{float(best['dVtotal_km_s']):.4f} km/s. The result is therefore a physically valid low-energy "
        f"end of the Pareto front, while the Lambert cases provide the fast-transfer end."
    )

    add_heading(doc, "Appendix A. Implementation Notes", 1)
    add_bullets(
        doc,
        [
            "The MATLAB code is self-contained and uses only standard MATLAB functions.",
            "The target GEO state is generated from ERA, longitude, and GEO radius to avoid table transcription ambiguity.",
            "Both prograde and retrograde Lambert candidates are tested, then Pareto filtered.",
            "The phasing family is tested over integer revolution counts and only valid orbits are retained.",
            "Figures and CSV tables are written to the results folder for reproducibility.",
            "The animation uses a uniform time grid so the speed is not distorted by adaptive ODE step spacing.",
        ],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
